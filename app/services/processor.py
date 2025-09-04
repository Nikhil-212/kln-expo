import os
from datetime import datetime
import re
import spacy
from jinja2 import Template
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph

# Load spaCy model for NLP processing
try:
    nlp = spacy.load("en_core_web_sm")
except OSError:
    import subprocess
    subprocess.run(["python", "-m", "spacy", "download", "en_core_web_sm"])
    nlp = spacy.load("en_core_web_sm")

class LegalDocumentProcessor:
    def __init__(self):
        self.document_types = {
            'rental_agreement': {
                'keywords': ['rental', 'rent', 'lease', 'tenant', 'landlord', 'monthly'],
                'required_fields': ['landlord', 'tenant', 'address', 'rent_amount', 'start_date', 'duration'],
                'template': 'rental_agreement_template.txt'
            },
            'land_sale_deed': {
                'keywords': ['sale', 'deed', 'property', 'buyer', 'seller', 'purchase'],
                'required_fields': ['seller', 'buyer', 'property_description', 'sale_amount'],
                'template': 'land_sale_deed_template.txt'
            },
            'power_of_attorney': {
                'keywords': ['power', 'attorney', 'delegate', 'authority', 'behalf'],
                'required_fields': ['grantor', 'attorney', 'powers', 'duration'],
                'template': 'power_of_attorney_template.txt'
            },
            'house_lease': {
                'keywords': ['house', 'lease', 'lessor', 'lessee', 'property'],
                'required_fields': ['lessor', 'lessee', 'property_address', 'lease_amount', 'start_date', 'duration'],
                'template': 'house_lease_template.txt'
            }
        }

        self.entity_patterns = {
            'amounts': r'(?:Rs\.?|INR)?\s*(\d+(?:,\d+)*(?:\.\d{2})?)\s*(?:rupees?|Rs\.?|INR)?',
            'dates': r'\d{1,2}[-/]\d{1,2}[-/]\d{4}|\d{1,2}(?:st|nd|rd|th)?\s+(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)\s+\d{4}',
            'locations': r'\b(?:at|in)\s+([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)',
            'durations': r'(\d+)\s*(?:year|month|week|day)s?'
        }

    def classify_document_type(self, prompt):
        """Classify the document type based on the user's prompt"""
        prompt_lower = prompt.lower()

        # Count keyword matches for each document type
        scores = {}
        for doc_type, info in self.document_types.items():
            score = sum(1 for keyword in info['keywords'] if keyword in prompt_lower)
            scores[doc_type] = score

        # Return the document type with the highest score
        if scores:
            return max(scores, key=scores.get)
        return None

    def extract_entities(self, prompt):
        """Extract entities from the user's prompt using spaCy and regex patterns"""
        doc = nlp(prompt)
        entities = {}

        # Extract named entities using spaCy
        for ent in doc.ents:
            if ent.label_ == 'PERSON':
                if 'landlord' not in entities:
                    entities['landlord'] = ent.text
                elif 'tenant' not in entities:
                    entities['tenant'] = ent.text
                elif 'seller' not in entities:
                    entities['seller'] = ent.text
                elif 'buyer' not in entities:
                    entities['buyer'] = ent.text
                elif 'grantor' not in entities:
                    entities['grantor'] = ent.text
                elif 'attorney' not in entities:
                    entities['attorney'] = ent.text
                elif 'lessor' not in entities:
                    entities['lessor'] = ent.text
                elif 'lessee' not in entities:
                    entities['lessee'] = ent.text
            elif ent.label_ == 'GPE' or ent.label_ == 'LOC':
                if 'address' not in entities:
                    entities['address'] = ent.text
                elif 'property_address' not in entities:
                    entities['property_address'] = ent.text

        # Extract amounts using regex
        amount_match = re.search(self.entity_patterns['amounts'], prompt, re.IGNORECASE)
        if amount_match:
            entities['rent_amount'] = amount_match.group(1)
            entities['sale_amount'] = amount_match.group(1)
            entities['lease_amount'] = amount_match.group(1)

        # Extract dates
        date_match = re.search(self.entity_patterns['dates'], prompt, re.IGNORECASE)
        if date_match:
            entities['start_date'] = date_match.group(0)
            entities['sale_date'] = date_match.group(0)

        # Extract durations
        duration_match = re.search(self.entity_patterns['durations'], prompt, re.IGNORECASE)
        if duration_match:
            entities['duration'] = duration_match.group(0)

        return entities

    def generate_document(self, doc_type, entities):
        """Generate document content by filling the template with extracted entities"""
        if doc_type not in self.document_types:
            raise ValueError(f"Unsupported document type: {doc_type}")

        template_path = os.path.join('templates', self.document_types[doc_type]['template'])

        try:
            with open(template_path, 'r', encoding='utf-8') as f:
                template_content = f.read()

            template = Template(template_content)

            # Fill in the entities, using defaults for missing fields
            filled_data = {}
            for field in self.document_types[doc_type]['required_fields']:
                filled_data[field] = entities.get(field, f'[{field.replace("_", " ").title()}]')

            # Add current date information
            now = datetime.now()
            filled_data.update({
                'date': now.strftime('%d'),
                'month': now.strftime('%B'),
                'year': now.strftime('%Y'),
                'execution_date': now.strftime('%d/%m/%Y'),
                'execution_place': 'City'
            })

            document = template.render(**filled_data)
            return document

        except FileNotFoundError:
            raise ValueError(f"Template file not found: {template_path}")
        except Exception as e:
            raise ValueError(f"Error generating document: {str(e)}")

    def generate_docx(self, content, filename):
        """Generate a .docx file from the document content"""
        doc = Document()
        doc.add_heading('Legal Document', 0)

        # Split content into paragraphs and add to document
        paragraphs = content.split('\n\n')
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para.strip())

        doc.save(filename)

    def generate_pdf(self, content, filename):
        """Generate a .pdf file from the document content"""
        doc = SimpleDocTemplate(filename, pagesize=letter)
        styles = getSampleStyleSheet()

        # Split content into paragraphs
        paragraphs = content.split('\n\n')
        story = []

        for para in paragraphs:
            if para.strip():
                story.append(Paragraph(para.strip(), styles['Normal']))

        doc.build(story)
