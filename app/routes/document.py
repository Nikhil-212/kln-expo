"""Document generation routes."""
from flask import render_template, request, flash, session, jsonify
from . import document_bp
from app.services.processor import LegalDocumentProcessor
from app.models.history import add_user_history
import spacy
import json
import tempfile
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_JUSTIFY
from reportlab.lib.units import inch
from datetime import datetime
import re

nlp = spacy.load('en_core_web_sm')
processor = LegalDocumentProcessor()

def get_default_data_for_document(doc_type, language):
    """Get default data for realistic document generation"""
    current_date = datetime.now()
    
    defaults = {
        'date': current_date.strftime('%B %d, %Y'),
        'month': current_date.strftime('%B'),
        'year': current_date.strftime('%Y'),
        'execution_date': current_date.strftime('%d/%m/%Y'),
        'execution_place': 'Chennai',
        'witness1_name': 'Mr. Witness One',
        'witness1_address': 'No. 123, Main Street, Chennai - 600001',
        'witness2_name': 'Mr. Witness Two', 
        'witness2_address': 'No. 456, Second Street, Chennai - 600002',
        'jurisdiction': 'Chennai',
        'registration_office': 'Chennai',
        'stamp_duty_bearer': 'Vendee'
    }
    
    if doc_type == 'rental_agreement':
        defaults.update({
            'owner_age': '45',
            'renter_age': '35',
            'owner_father': 'Father Name',
            'renter_father': 'Father Name',
            'owner_address': 'No. 123, Main Street',
            'owner_city': 'Chennai',
            'owner_pincode': '600001',
            'renter_address': 'No. 456, Second Street',
            'renter_city': 'Chennai',
            'renter_pincode': '600002',
            'property_address': 'No. 789, Property Street',
            'property_city': 'Chennai',
            'property_pincode': '600073',
            'start_date': '1st April 2024',
            'effective_date': '1st April 2024',
            'duration': '11',
            'renewal_period': '11',
            'rent_amount': '15,000',
            'rent_amount_words': 'Fifteen Thousand',
            'rent_due_date': '1st',
            'rent_increase_percentage': '10',
            'security_deposit': '30,000',
            'security_deposit_words': 'Thirty Thousand',
            'notice_period': '2'
        })
    elif doc_type == 'land_sale_deed':
        defaults.update({
            'seller_age': '50',
            'buyer_age': '40',
            'seller_father': 'Father Name',
            'buyer_father': 'Father Name',
            'seller_address': 'No. 123, Main Street',
            'seller_city': 'Chennai',
            'seller_pincode': '600001',
            'buyer_address': 'No. 456, Second Street',
            'buyer_city': 'Chennai',
            'buyer_pincode': '600002',
            'property_address': 'No. 789, Property Street',
            'property_city': 'Chennai',
            'property_pincode': '600073',
            'sale_amount': '50,00,000',
            'sale_amount_words': 'Fifty Lakhs',
            'sale_date': '1st April 2024',
            'survey_number': '123/45',
            'area': '2400',
            'north_boundary': 'Main Road',
            'south_boundary': 'Residential Area',
            'east_boundary': 'Park',
            'west_boundary': 'Commercial Area'
        })
    elif doc_type == 'power_of_attorney':
        defaults.update({
            'principal_age': '55',
            'attorney_age': '40',
            'principal_father': 'Father Name',
            'attorney_father': 'Father Name',
            'principal_address': 'No. 123, Main Street',
            'principal_city': 'Chennai',
            'principal_pincode': '600001',
            'attorney_address': 'No. 456, Second Street',
            'attorney_city': 'Chennai',
            'attorney_pincode': '600002',
            'matter_description': 'property management and legal representation',
            'effective_date': '1st April 2024',
            'expiry_date': '31st March 2025'
        })
    elif doc_type == 'house_lease':
        defaults.update({
            'lessor_age': '50',
            'lessee_age': '35',
            'lessor_father': 'Father Name',
            'lessee_father': 'Father Name',
            'lessor_address': 'No. 123, Main Street',
            'lessor_city': 'Chennai',
            'lessor_pincode': '600001',
            'lessee_address': 'No. 456, Second Street',
            'lessee_city': 'Chennai',
            'lessee_pincode': '600002',
            'property_address': 'No. 789, Property Street',
            'property_city': 'Chennai',
            'property_pincode': '600073',
            'lease_period': '2',
            'start_date': '1st April 2024',
            'end_date': '31st March 2026',
            'lease_amount': '25,000',
            'lease_amount_words': 'Twenty Five Thousand',
            'rent_due_date': '1st',
            'security_deposit': '50,000',
            'security_deposit_words': 'Fifty Thousand',
            'notice_period': '3',
            'number_of_rooms': '3'
        })
    
    # Apply language-specific translations
    translations = {
        'en': {
            'Father Name': 'Father Name',
            'Mr. Witness One': 'Mr. Witness One',
            'Chennai': 'Chennai'
        },
        'hi': {
            'Father Name': 'अपने पिता का नाम',
            'Mr. Witness One': 'श्री विजय एक',
            'Chennai': 'चेन्नई'
        },
        'bn': {
            'Father Name': 'আমার পিতার নাম',
            'Mr. Witness One': 'শ্রী বিজয় এক',
            'Chennai': 'চেন্নাই'
        },
        'te': {
            'Father Name': 'నా పిల్లి పేరు',
            'Mr. Witness One': 'శ్రీ విజయ ఒక',
            'Chennai': 'చెన్నై'
        },
        'mr': {
            'Father Name': 'माझा वडील यांचा नाव',
            'Mr. Witness One': 'श्री విజయ एक',
            'Chennai': 'चेन्नई'
        },
        'ur': {
            'Father Name': 'میرے والد کا نام',
            'Mr. Witness One': 'شری ویجی ہے',
            'Chennai': 'چینనాی'
        },
        'gu': {
            'Father Name': 'માઝા પિતાનું નામ',
            'Mr. Witness One': 'શ્રી વિજય એક',
            'Chennai': 'ચેન્નઈ'
        },
        'kn': {
            'Father Name': 'ನನ್ನ ಪಿತಾನ ಹೆಸರು',
            'Mr. Witness One': 'ಶ್ರೀ ವಿಜಯ ಒಂದು',
            'Chennai': 'ಚೆನ್ನಾಯ'
        },
        'or': {
            'Father Name': 'ଆମଦ୍ବାରା ପିତାଙ୍କ ନାମ',
            'Mr. Witness One': 'ଶ୍ରୀ ବିଜଯ଼ ଏକ',
            'Chennai': 'ଚେନ୍ନାଇ'
        },
        'ta': {
            'Father Name': 'என் தந்தை பெயர்',
            'Mr. Witness One': 'சிறுவர் விஜய் ஒன்று',
            'Chennai': 'சென்னை'
        }
    }

    for key, value in defaults.items():
        if isinstance(value, str) and value in translations.get(language, {}):
            defaults[key] = translations[language][value]
        elif isinstance(value, str) and value.lower() in translations.get(language, {}):
            defaults[key] = translations[language][value.lower()]

    return defaults

documents = {
    'rental_agreement': {
        'templates': {
            'en': 'rental_agreement_template.txt',
            'hi': 'hi/rental_agreement_template.txt',
            'bn': 'bn/rental_agreement_template.txt',
            'te': 'te/rental_agreement_template.txt',
            'mr': 'mr/rental_agreement_template.txt',
            'ur': 'ur/rental_agreement_template.txt',
            'gu': 'gu/rental_agreement_template.txt',
            'kn': 'kn/rental_agreement_template.txt',
            'or': 'or/rental_agreement_template.txt',
            'ta': 'ta/rental_agreement_template.txt'
        },
        'fields': ['owner_name', 'owner_age', 'owner_father', 'owner_address', 'owner_city', 'owner_pincode',
                   'renter_name', 'renter_age', 'renter_father', 'renter_address', 'renter_city', 'renter_pincode',
                   'property_address', 'property_city', 'property_pincode', 'start_date', 'effective_date',
                   'duration', 'renewal_period', 'rent_amount', 'rent_amount_words', 'rent_due_date',
                   'rent_increase_percentage', 'security_deposit', 'security_deposit_words', 'notice_period',
                   'jurisdiction', 'witness1_name', 'witness1_address', 'witness2_name', 'witness2_address',
                   'execution_date', 'execution_place', 'date', 'month', 'year']
    },
    'land_sale_deed': {
        'templates': {
            'en': 'land_sale_deed_template.txt',
            'hi': 'hi/land_sale_deed_template.txt',
            'bn': 'bn/land_sale_deed_template.txt',
            'te': 'te/land_sale_deed_template.txt',
            'mr': 'mr/land_sale_deed_template.txt',
            'ur': 'ur/land_sale_deed_template.txt',
            'gu': 'gu/land_sale_deed_template.txt',
            'kn': 'kn/land_sale_deed_template.txt',
            'or': 'or/land_sale_deed_template.txt',
            'ta': 'ta/land_sale_deed_template.txt'
        },
        'fields': ['seller', 'seller_age', 'seller_father', 'seller_address', 'seller_city', 'seller_pincode',
                   'buyer', 'buyer_age', 'buyer_father', 'buyer_address', 'buyer_city', 'buyer_pincode',
                   'sale_amount', 'sale_amount_words', 'stamp_duty_bearer', 'registration_office',
                   'survey_number', 'area', 'north_boundary', 'south_boundary', 'east_boundary', 'west_boundary',
                   'property_address', 'property_city', 'property_pincode', 'witness1_name', 'witness1_address',
                   'witness2_name', 'witness2_address', 'execution_date', 'execution_place', 'date', 'month', 'year']
    },
    'power_of_attorney': {
        'templates': {
            'en': 'power_of_attorney_template.txt',
            'hi': 'hi/power_of_attorney_template.txt',
            'bn': 'bn/power_of_attorney_template.txt',
            'te': 'te/power_of_attorney_template.txt',
            'mr': 'mr/power_of_attorney_template.txt',
            'ur': 'ur/power_of_attorney_template.txt',
            'gu': 'gu/power_of_attorney_template.txt',
            'kn': 'kn/power_of_attorney_template.txt',
            'or': 'or/power_of_attorney_template.txt',
            'ta': 'ta/power_of_attorney_template.txt'
        },
        'fields': ['principal', 'principal_age', 'principal_father', 'principal_address', 'principal_city', 'principal_pincode',
                   'attorney', 'attorney_age', 'attorney_father', 'attorney_address', 'attorney_city', 'attorney_pincode',
                   'matter_description', 'effective_date', 'expiry_date', 'registration_office', 'stamp_duty_bearer',
                   'witness1_name', 'witness1_address', 'witness2_name', 'witness2_address', 'execution_date', 'execution_place',
                   'date', 'month', 'year']
    },
    'house_lease': {
        'templates': {
            'en': 'house_lease_template.txt',
            'hi': 'hi/house_lease_template.txt',
            'bn': 'bn/house_lease_template.txt',
            'te': 'te/house_lease_template.txt',
            'mr': 'mr/house_lease_template.txt',
            'ur': 'ur/house_lease_template.txt',
            'gu': 'gu/house_lease_template.txt',
            'kn': 'kn/house_lease_template.txt',
            'or': 'or/house_lease_template.txt',
            'ta': 'ta/house_lease_template.txt'
        },
        'fields': ['lessor', 'lessor_age', 'lessor_father', 'lessor_address', 'lessor_city', 'lessor_pincode',
                   'lessee', 'lessee_age', 'lessee_father', 'lessee_address', 'lessee_city', 'lessee_pincode',
                   'property_address', 'property_city', 'property_pincode', 'lease_period', 'start_date', 'end_date',
                   'lease_amount', 'lease_amount_words', 'rent_due_date', 'security_deposit', 'security_deposit_words',
                   'notice_period', 'jurisdiction', 'number_of_rooms', 'witness1_name', 'witness1_address',
                   'witness2_name', 'witness2_address', 'execution_date', 'execution_place', 'date', 'month', 'year']
    }
}

@document_bp.route('/document/<doc_type>')
def document_form(doc_type):
    if doc_type not in documents:
        return render_template('index.html', error='Invalid document type')
    languages = list(documents[doc_type]['templates'].keys())
    return render_template('document_form.html', doc_type=doc_type, fields=documents[doc_type]['fields'], languages=languages)

@document_bp.route('/generate', methods=['POST'])
def generate_document():
    doc_type = request.form.get('doc_type')
    language = request.form.get('language', 'en')
    if not doc_type or doc_type not in documents:
        return render_template('index.html', error='Invalid document type')

    # Collect field values and check for missing ones
    fields = documents[doc_type]['fields']
    data = {field: request.form.get(field, '').strip() for field in fields}
    # No longer require all fields to be filled. Missing fields will simply be empty in the template.
    # missing_fields = [field for field, value in data.items() if not value]

    # if missing_fields:
    #     error_msg = f"Please fill in all required fields: {', '.join(missing_fields)}"
    #     languages = list(documents[doc_type]['templates'].keys())
    #     return render_template('document_form.html', doc_type=doc_type, fields=fields, error=error_msg, values=data, languages=languages, selected_language=language)

    try:
        document = processor.generate_document(doc_type, data, language)
        # Use spaCy to extract entities
        doc_nlp = nlp(document)
        entities = [(ent.text, ent.label_) for ent in doc_nlp.ents]

        # Log history
        if 'user_id' in session:
            add_user_history(session['user_id'], 'generate_document', f'Generated {doc_type}')

        return render_template('view_document.html', doc_type=doc_type, content=document, entities=entities)
    except Exception as e:
        import traceback
        error_message = f"Error generating document: {str(e)}"
        print(traceback.format_exc())
        flash(error_message, 'danger')
        languages = list(documents[doc_type]['templates'].keys())
        return render_template('document_form.html', doc_type=doc_type, fields=fields, error=error_message, values=data, languages=languages, selected_language=language)

@document_bp.route('/generate_from_prompt', methods=['POST'])
def generate_from_prompt():
    prompt = request.form.get('prompt', '').strip()
    if not prompt:
        flash('Please enter a prompt describing the document you want to generate.', 'error')
        return render_template('index.html')

    try:
        # Extract language from prompt, default to 'en'
        language_match = re.search(r'in\s+([a-zA-Z]+)', prompt)
        language = language_match.group(1).lower() if language_match else 'en'

        # Classify document type
        doc_type = processor.classify_document_type(prompt)
        if not doc_type:
            flash('Could not determine document type from your prompt. Please try rephrasing.', 'error')
            return render_template('index.html')

        # Extract entities
        entities = processor.extract_entities(prompt)

        # Generate document
        document = processor.generate_document(doc_type, entities, language=language)

        # Extract entities from generated document for display
        doc_nlp = nlp(document)
        extracted_entities = [(ent.text, ent.label_) for ent in doc_nlp.ents]

        # Log history
        if 'user_id' in session:
            add_user_history(session['user_id'], 'generate_from_prompt', f'Generated {doc_type} from prompt')

        flash(f'Document type classified as: {doc_type.replace("_", " ").title()}', 'success')
        return render_template('view_document.html', doc_type=doc_type, content=document, entities=extracted_entities, prompt=prompt)

    except Exception as e:
        flash(f'Error generating document: {str(e)}', 'error')
        return render_template('index.html')

@document_bp.route('/download/<doc_type>/<format>')
def download_document(doc_type, format):
    # This would need the document content stored in session or database
    # For now, return a placeholder
    flash('Download functionality will be implemented with proper session management.', 'info')
    return render_template('index.html')

@document_bp.route('/api/process-prompt', methods=['POST'])
def api_process_prompt():
    """Process user prompt and extract information"""
    try:
        data = request.get_json()
        prompt = data.get('prompt', '')
        
        if not prompt:
            return jsonify({'error': 'No prompt provided'}), 400
        
        # Classify document type
        doc_type = processor.classify_document_type(prompt)
        
        # Extract entities
        entities = processor.extract_entities(prompt)
        
        # Identify missing fields
        missing_fields = processor.identify_missing_fields(doc_type, entities)
        
        # Prepare response
        response = {
            'document_type': doc_type,
            'extracted_entities': entities,
            'missing_fields': missing_fields,
            'status': 'success'
        }
        
        return jsonify(response)
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@document_bp.route('/api/generate-document', methods=['POST'])
def api_generate_document():
    """Generate final document with all required data"""
    try:
        data = request.get_json()
        doc_type = data.get('document_type')
        filled_data = data.get('filled_data', {})
        format_type = data.get('format', 'docx')
        language = data.get('language', 'en')
        custom_template = data.get('custom_template')
        
        if not doc_type or not filled_data:
            return jsonify({'error': 'Missing required data'}), 400
        
        # Generate document content
        if custom_template:
            from jinja2 import Template
            try:
                template = Template(custom_template)
                document_content = template.render(**filled_data)
            except Exception as e:
                return jsonify({'error': f'Error rendering custom template: {str(e)}'}), 400
        else:
            # Merge filled_data with default values for a complete document
            complete_data = get_default_data_for_document(doc_type, language)
            complete_data.update(filled_data)  # User data overrides defaults
            document_content = processor.generate_document(doc_type, complete_data, language=language)
        
        # Create file based on format
        if format_type == 'docx':
            return create_docx_file(document_content, doc_type)
        elif format_type == 'pdf':
            return create_pdf_file(document_content, doc_type)
        else:
            return jsonify({'error': 'Unsupported format'}), 400
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500

def create_docx_file(content, doc_type):
    """Create DOCX file from content"""
    try:
        doc = Document()
        doc.add_heading(f'{doc_type.replace("_", " ").title()}', 0)
        
        # Add content
        paragraphs = content.split('\n')
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para.strip())
        
        # Save to temporary file
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        temp_file.close()
        
        from flask import send_file
        from datetime import datetime
        return send_file(
            temp_file.name,
            as_attachment=True,
            download_name=f'{doc_type}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    
    except Exception as e:
        return jsonify({'error': f'Error creating DOCX: {str(e)}'}), 500

def create_pdf_file(content, doc_type):
    """Create PDF file from content"""
    try:
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
        temp_path = temp_file.name
        temp_file.close()

        doc = SimpleDocTemplate(temp_path, pagesize=A4, leftMargin=54, rightMargin=54, topMargin=54, bottomMargin=54)
        styles = getSampleStyleSheet()

        body_style = ParagraphStyle(
            name='Body',
            parent=styles['Normal'],
            fontSize=11,
            leading=16,
            alignment=TA_JUSTIFY,
        )
        title_style = ParagraphStyle(
            name='Title',
            parent=styles['Heading1'],
            fontSize=18,
            leading=22,
            spaceAfter=12,
        )

        story = []
        story.append(Paragraph(doc_type.replace('_', ' ').title(), title_style))
        story.append(Spacer(1, 0.2 * inch))

        # Convert plain text line breaks to simple paragraphs
        for block in content.split('\n\n'):
            block_html = block.strip().replace('\n', '<br/>')
            if not block_html:
                continue
            story.append(Paragraph(block_html, body_style))
            story.append(Spacer(1, 0.12 * inch))

        doc.build(story)

        from flask import send_file
        from datetime import datetime
        return send_file(
            temp_path,
            as_attachment=True,
            download_name=f'{doc_type}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.pdf',
            mimetype='application/pdf'
        )
    except Exception as e:
        return jsonify({'error': f'Error creating PDF: {str(e)}'}), 500

@document_bp.route('/edit_document', methods=['GET', 'POST'])
def edit_document():
    doc_type = request.args.get('doc_type')
    content = request.args.get('content')

    if request.method == 'POST':
        edited_content = request.form.get('edited_content')
        doc_type = request.form.get('doc_type')
        return render_template('view_document.html', doc_type=doc_type, content=edited_content, entities=[]) # Simplified entities for now

    return render_template('edit_document.html', doc_type=doc_type, content=content)
